"""
TRINETRA — Word (.docx) Report Generator

Converts the AI chatbot's markdown-formatted SOC report into a polished,
professional Word document that can be sent to a SOC analyst or client.
"""

import re
from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x8B, 0x5C, 0xF6)
DARK = RGBColor(0x1a, 0x1a, 0x2e)
MUTED = RGBColor(0x6b, 0x72, 0x80)
CODE_COLOR = RGBColor(0x6d, 0x28, 0xd9)

BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_page_border(doc):
    sectPr = doc.sections[0]._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), "8B5CF6")
        pgBorders.append(el)
    sectPr.append(pgBorders)


def _set_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK


def _add_cover_page(doc, title, target, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRINETRA")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle.upper())
    run2.font.size = Pt(13)
    run2.font.color.rgb = MUTED
    run2.font.bold = True

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(title)
    run3.font.size = Pt(20)
    run3.font.bold = True

    doc.add_paragraph()

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows_data = [
        ("Target", target or "-"),
        ("Report Date", datetime.now().strftime("%d %B %Y, %H:%M")),
        ("Prepared By", "TRINETRA OSINT Platform v1.0 - AI Assistant"),
    ]
    for row, (label, value) in zip(meta_table.rows, rows_data):
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = MUTED
        row.cells[1].text = value

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer.add_run(
        "This report was generated from open-source intelligence (OSINT) data. "
        "For authorized use only."
    )
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = MUTED
    frun.font.italic = True

    doc.add_page_break()


def _parse_markdown_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _add_inline(paragraph, text):
    """Add text with **bold** and `inline code` support to a paragraph."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_COLOR
        else:
            paragraph.add_run(part)


def _strip_preamble(lines):
    """Drop any chatty intro text before the first markdown heading, and drop
    a redundant top-level (H1) title line since the cover page already shows it."""
    first_heading_idx = next((i for i, l in enumerate(lines) if l.strip().startswith("#")), None)
    if first_heading_idx is None:
        return lines
    lines = lines[first_heading_idx:]
    if re.match(r"^#\s+\S", lines[0].strip()) and not lines[0].strip().startswith("##"):
        lines = lines[1:]
    return lines


def markdown_to_docx(doc, markdown_text):
    lines = markdown_text.replace("\r\n", "\n").split("\n")
    lines = _strip_preamble(lines)
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule dividers (---, ___, ***) — skip, don't render as text
        if re.match(r"^(-{3,}|_{3,}|\*{3,})$", stripped):
            i += 1
            continue

        # Tables
        if stripped.startswith("|"):
            rows, i = _parse_markdown_table(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < len(table.columns):
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            cp = cell.paragraphs[0]
                            _add_inline(cp, cell_text)
                            if r_idx == 0:
                                for run in cp.runs:
                                    run.font.bold = True
                                _shade_cell(cell, "EDE9FE")
                doc.add_paragraph()
            continue

        # Headers
        if stripped.startswith("#### "):
            h = doc.add_heading(level=4)
            _add_inline(h, stripped[5:])
            i += 1
            continue
        if stripped.startswith("### "):
            h = doc.add_heading(level=3)
            _add_inline(h, stripped[4:])
            i += 1
            continue
        if stripped.startswith("## "):
            h = doc.add_heading(level=2)
            _add_inline(h, stripped[3:])
            i += 1
            continue
        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            _add_inline(h, stripped[2:])
            i += 1
            continue

        # Indentation depth (2 spaces = 1 nesting level, capped at level 3)
        indent = len(line) - len(line.lstrip(" "))
        level = min(indent // 2, 2)

        # Bullets
        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style=BULLET_STYLES[level])
            _add_inline(p, stripped[2:])
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style=NUMBER_STYLES[level])
            _add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1


def build_report_docx(title, target, markdown_text, subtitle="SOC Investigation Report"):
    doc = Document()
    _set_base_styles(doc)
    _add_page_border(doc)
    _add_cover_page(doc, title, target, subtitle)
    markdown_to_docx(doc, markdown_text)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf